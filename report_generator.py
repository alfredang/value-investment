"""
Report Generator for Value Investment Analysis.

Generates comprehensive investment analysis reports and exports to Word/PDF format.
"""
import io
from datetime import datetime
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
import pandas as pd

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class CompanyAnalysis:
    """Analysis data for a single company."""
    symbol: str
    company_name: str
    sector: str
    market: str

    # Fundamentals
    gross_margin: Optional[float] = None
    net_margin: Optional[float] = None
    roe: Optional[float] = None
    roa: Optional[float] = None
    debt_to_equity: Optional[float] = None
    fcf_margin: Optional[float] = None
    roic_wacc: Optional[float] = None
    rote_wacc: Optional[float] = None
    revenue_growth_5y: Optional[float] = None
    eps_growth_5y: Optional[float] = None

    # Valuation
    market_cap: Optional[float] = None
    epv: Optional[float] = None
    epv_mc_ratio: Optional[float] = None
    valuation_status: str = "N/A"

    # Anomaly Detection
    risk_level: str = "N/A"
    m_score: Optional[float] = None
    z_score: Optional[float] = None
    f_score: Optional[float] = None
    sloan_ratio: Optional[float] = None
    anomalies: List[Dict] = None

    # AI Analysis
    ai_summary: str = ""
    ai_recommendation: str = ""


class ReportGenerator:
    """Generate comprehensive investment analysis reports."""

    def __init__(self):
        self.report_date = datetime.now().strftime("%Y-%m-%d")

    def generate_markdown_report(
        self,
        companies: List[CompanyAnalysis],
        screening_criteria: Dict[str, Any],
        title: str = "Value Investment Analysis Report"
    ) -> str:
        """
        Generate a comprehensive Markdown report.

        Args:
            companies: List of CompanyAnalysis objects
            screening_criteria: The criteria used for screening
            title: Report title

        Returns:
            Markdown formatted report string
        """
        lines = []

        # Title and metadata
        lines.append(f"# {title}")
        lines.append(f"\n**Generated:** {self.report_date}")
        lines.append(f"**Companies Analyzed:** {len(companies)}")
        lines.append("")

        # Executive Summary
        lines.append("## Executive Summary")
        lines.append("")
        undervalued = [c for c in companies if c.valuation_status == "Undervalued"]
        high_risk = [c for c in companies if c.risk_level in ["HIGH RISK", "ELEVATED RISK"]]

        lines.append(f"- **Total Companies Screened:** {len(companies)}")
        lines.append(f"- **Undervalued Opportunities:** {len(undervalued)}")
        lines.append(f"- **High/Elevated Risk:** {len(high_risk)}")
        lines.append("")

        # Screening Criteria Used
        lines.append("## Screening Criteria")
        lines.append("")
        lines.append("| Criterion | Threshold |")
        lines.append("|-----------|-----------|")
        for key, value in screening_criteria.items():
            formatted_key = key.replace("_", " ").title()
            lines.append(f"| {formatted_key} | {value} |")
        lines.append("")

        # Summary Table
        lines.append("## Company Overview")
        lines.append("")
        lines.append("| Symbol | Company | Sector | Valuation | Risk Level | EPV/MC |")
        lines.append("|--------|---------|--------|-----------|------------|--------|")
        for c in companies:
            epv_mc = f"{c.epv_mc_ratio:.2f}" if c.epv_mc_ratio else "N/A"
            lines.append(f"| {c.symbol} | {c.company_name[:30]} | {c.sector[:20]} | {c.valuation_status} | {c.risk_level} | {epv_mc} |")
        lines.append("")

        # Detailed Analysis for Each Company
        lines.append("## Detailed Company Analysis")
        lines.append("")

        for company in companies:
            lines.extend(self._generate_company_section(company))
            lines.append("")

        # Recommendations Summary
        lines.append("## Investment Recommendations")
        lines.append("")

        # Group by recommendation
        strong_buys = [c for c in companies if c.valuation_status == "Undervalued" and c.risk_level in ["LOW RISK", "MINIMAL RISK", "MODERATE RISK"]]
        cautious = [c for c in companies if c.valuation_status == "Undervalued" and c.risk_level in ["HIGH RISK", "ELEVATED RISK"]]

        if strong_buys:
            lines.append("### Strong Buy Candidates")
            lines.append("*Undervalued with acceptable risk levels*")
            lines.append("")
            for c in strong_buys:
                lines.append(f"- **{c.symbol}** ({c.company_name}): {c.valuation_status}, {c.risk_level}")
            lines.append("")

        if cautious:
            lines.append("### Requires Further Investigation")
            lines.append("*Undervalued but elevated risk - investigate anomalies before investing*")
            lines.append("")
            for c in cautious:
                lines.append(f"- **{c.symbol}** ({c.company_name}): {c.valuation_status}, {c.risk_level}")
            lines.append("")

        # Disclaimer
        lines.append("---")
        lines.append("")
        lines.append("## Disclaimer")
        lines.append("")
        lines.append("*This report is generated for informational purposes only and does not constitute financial advice. ")
        lines.append("Always conduct your own due diligence and consult with a qualified financial advisor before making investment decisions.*")
        lines.append("")
        lines.append(f"*Report generated on {self.report_date} using Value Investment Analysis Tool*")

        return "\n".join(lines)

    def _generate_company_section(self, company: CompanyAnalysis) -> List[str]:
        """Generate detailed section for a single company."""
        lines = []

        lines.append(f"### {company.symbol} - {company.company_name}")
        lines.append("")
        lines.append(f"**Sector:** {company.sector} | **Market:** {company.market}")
        lines.append("")

        # Valuation Box
        lines.append("#### Valuation Assessment")
        lines.append("")
        epv_mc = f"{company.epv_mc_ratio:.2f}" if company.epv_mc_ratio else "N/A"
        lines.append(f"| Metric | Value |")
        lines.append(f"|--------|-------|")
        lines.append(f"| **Valuation Status** | **{company.valuation_status}** |")
        lines.append(f"| Market Cap | ${company.market_cap:.1f}M |" if company.market_cap else "| Market Cap | N/A |")
        lines.append(f"| EPV | ${company.epv:.1f}M |" if company.epv else "| EPV | N/A |")
        lines.append(f"| EPV/MC Ratio | {epv_mc} |")
        lines.append("")

        # Fundamentals
        lines.append("#### Key Fundamentals")
        lines.append("")
        lines.append("| Metric | Value | Status |")
        lines.append("|--------|-------|--------|")

        def fmt(val, suffix="%", threshold=None, higher_better=True):
            if val is None:
                return "N/A", ""
            formatted = f"{val:.1f}{suffix}"
            if threshold is not None:
                if higher_better:
                    status = "Good" if val >= threshold else "Below Target"
                else:
                    status = "Good" if val <= threshold else "Above Target"
                return formatted, status
            return formatted, ""

        gm, gm_s = fmt(company.gross_margin, "%", 20, True)
        nm, nm_s = fmt(company.net_margin, "%", 5, True)
        roe, roe_s = fmt(company.roe, "%", 10, True)
        roa, roa_s = fmt(company.roa, "%", 5, True)
        de, de_s = fmt(company.debt_to_equity, "", 1.5, False)
        fcf, fcf_s = fmt(company.fcf_margin, "%", 5, True)
        roic, roic_s = fmt(company.roic_wacc, "", 0, True)

        lines.append(f"| Gross Margin | {gm} | {gm_s} |")
        lines.append(f"| Net Margin | {nm} | {nm_s} |")
        lines.append(f"| ROE | {roe} | {roe_s} |")
        lines.append(f"| ROA | {roa} | {roa_s} |")
        lines.append(f"| Debt/Equity | {de} | {de_s} |")
        lines.append(f"| FCF Margin | {fcf} | {fcf_s} |")
        lines.append(f"| ROIC-WACC | {roic} | {roic_s} |")
        lines.append("")

        # Risk Assessment
        lines.append("#### Risk Assessment")
        lines.append("")
        lines.append(f"**Overall Risk Level: {company.risk_level}**")
        lines.append("")
        lines.append("| Quality Score | Value | Alert |")
        lines.append("|---------------|-------|-------|")

        m_alert = "ALERT" if company.m_score and company.m_score > -1.78 else "OK"
        z_alert = "ALERT" if company.z_score and company.z_score < 1.8 else "OK"
        f_alert = "ALERT" if company.f_score and company.f_score < 3 else "OK"
        s_alert = "ALERT" if company.sloan_ratio and abs(company.sloan_ratio) > 10 else "OK"

        lines.append(f"| M-Score | {company.m_score:.2f if company.m_score else 'N/A'} | {m_alert} |")
        lines.append(f"| Z-Score | {company.z_score:.2f if company.z_score else 'N/A'} | {z_alert} |")
        lines.append(f"| F-Score | {int(company.f_score) if company.f_score else 'N/A'} | {f_alert} |")
        sloan_val = f"{company.sloan_ratio:.1f}%" if company.sloan_ratio else "N/A"
        lines.append(f"| Sloan Ratio | {sloan_val} | {s_alert} |")
        lines.append("")

        # Anomalies
        if company.anomalies:
            lines.append("#### Detected Anomalies")
            lines.append("")
            high_anomalies = [a for a in company.anomalies if a.get('severity') == 'HIGH']
            if high_anomalies:
                lines.append("**High Severity Issues:**")
                for a in high_anomalies[:5]:
                    lines.append(f"- [{a.get('category', 'Unknown')}] {a.get('description', '')}")
                lines.append("")

        # AI Summary
        if company.ai_summary:
            lines.append("#### AI Analysis")
            lines.append("")
            lines.append(company.ai_summary)
            lines.append("")

        if company.ai_recommendation:
            lines.append("#### Recommendation")
            lines.append("")
            lines.append(company.ai_recommendation)
            lines.append("")

        lines.append("---")

        return lines

    def generate_word_document(
        self,
        companies: List[CompanyAnalysis],
        screening_criteria: Dict[str, Any],
        title: str = "Value Investment Analysis Report"
    ) -> io.BytesIO:
        """
        Generate a Word document report.

        Args:
            companies: List of CompanyAnalysis objects
            screening_criteria: The criteria used for screening
            title: Report title

        Returns:
            BytesIO buffer containing the Word document
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export. Install with: pip install python-docx")

        doc = Document()

        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Generated: {self.report_date}\n").bold = True
        meta.add_run(f"Companies Analyzed: {len(companies)}")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)

        undervalued = [c for c in companies if c.valuation_status == "Undervalued"]
        high_risk = [c for c in companies if c.risk_level in ["HIGH RISK", "ELEVATED RISK"]]

        summary = doc.add_paragraph()
        summary.add_run("Key Findings:\n").bold = True
        summary.add_run(f"  Total Companies Screened: {len(companies)}\n")
        summary.add_run(f"  Undervalued Opportunities: {len(undervalued)}\n")
        summary.add_run(f"  High/Elevated Risk: {len(high_risk)}\n")

        # Screening Criteria
        doc.add_heading("Screening Criteria", level=1)

        criteria_table = doc.add_table(rows=1, cols=2)
        criteria_table.style = 'Table Grid'
        hdr = criteria_table.rows[0].cells
        hdr[0].text = 'Criterion'
        hdr[1].text = 'Threshold'

        for key, value in screening_criteria.items():
            row = criteria_table.add_row().cells
            row[0].text = key.replace("_", " ").title()
            row[1].text = str(value)

        doc.add_paragraph()

        # Company Overview Table
        doc.add_heading("Company Overview", level=1)

        overview_table = doc.add_table(rows=1, cols=5)
        overview_table.style = 'Table Grid'
        hdr = overview_table.rows[0].cells
        hdr[0].text = 'Symbol'
        hdr[1].text = 'Company'
        hdr[2].text = 'Valuation'
        hdr[3].text = 'Risk Level'
        hdr[4].text = 'EPV/MC'

        for c in companies:
            row = overview_table.add_row().cells
            row[0].text = c.symbol
            row[1].text = c.company_name[:25]
            row[2].text = c.valuation_status
            row[3].text = c.risk_level
            row[4].text = f"{c.epv_mc_ratio:.2f}" if c.epv_mc_ratio else "N/A"

        doc.add_paragraph()

        # Detailed Analysis
        doc.add_heading("Detailed Company Analysis", level=1)

        for company in companies:
            self._add_company_to_doc(doc, company)

        # Recommendations
        doc.add_heading("Investment Recommendations", level=1)

        strong_buys = [c for c in companies if c.valuation_status == "Undervalued" and c.risk_level in ["LOW RISK", "MINIMAL RISK", "MODERATE RISK"]]
        cautious = [c for c in companies if c.valuation_status == "Undervalued" and c.risk_level in ["HIGH RISK", "ELEVATED RISK"]]

        if strong_buys:
            doc.add_heading("Strong Buy Candidates", level=2)
            para = doc.add_paragraph()
            para.add_run("Undervalued with acceptable risk levels:\n").italic = True
            for c in strong_buys:
                para.add_run(f"  {c.symbol} ({c.company_name})\n")

        if cautious:
            doc.add_heading("Requires Further Investigation", level=2)
            para = doc.add_paragraph()
            para.add_run("Undervalued but elevated risk - investigate before investing:\n").italic = True
            for c in cautious:
                para.add_run(f"  {c.symbol} ({c.company_name})\n")

        # Disclaimer
        doc.add_page_break()
        doc.add_heading("Disclaimer", level=1)
        disclaimer = doc.add_paragraph()
        disclaimer.add_run(
            "This report is generated for informational purposes only and does not constitute "
            "financial advice. Always conduct your own due diligence and consult with a qualified "
            "financial advisor before making investment decisions."
        ).italic = True

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def _add_company_to_doc(self, doc: Document, company: CompanyAnalysis):
        """Add a company section to the Word document."""
        doc.add_heading(f"{company.symbol} - {company.company_name}", level=2)

        info = doc.add_paragraph()
        info.add_run(f"Sector: ").bold = True
        info.add_run(f"{company.sector}  |  ")
        info.add_run(f"Market: ").bold = True
        info.add_run(f"{company.market}")

        # Valuation
        doc.add_heading("Valuation Assessment", level=3)
        val_table = doc.add_table(rows=4, cols=2)
        val_table.style = 'Table Grid'

        val_table.rows[0].cells[0].text = "Valuation Status"
        val_table.rows[0].cells[1].text = company.valuation_status
        val_table.rows[1].cells[0].text = "Market Cap"
        val_table.rows[1].cells[1].text = f"${company.market_cap:.1f}M" if company.market_cap else "N/A"
        val_table.rows[2].cells[0].text = "EPV"
        val_table.rows[2].cells[1].text = f"${company.epv:.1f}M" if company.epv else "N/A"
        val_table.rows[3].cells[0].text = "EPV/MC Ratio"
        val_table.rows[3].cells[1].text = f"{company.epv_mc_ratio:.2f}" if company.epv_mc_ratio else "N/A"

        # Risk
        doc.add_heading("Risk Assessment", level=3)
        risk_para = doc.add_paragraph()
        risk_para.add_run(f"Overall Risk Level: {company.risk_level}").bold = True

        # Quality Scores
        score_table = doc.add_table(rows=5, cols=3)
        score_table.style = 'Table Grid'
        score_table.rows[0].cells[0].text = "Score"
        score_table.rows[0].cells[1].text = "Value"
        score_table.rows[0].cells[2].text = "Status"

        score_table.rows[1].cells[0].text = "M-Score"
        score_table.rows[1].cells[1].text = f"{company.m_score:.2f}" if company.m_score else "N/A"
        score_table.rows[1].cells[2].text = "ALERT" if company.m_score and company.m_score > -1.78 else "OK"

        score_table.rows[2].cells[0].text = "Z-Score"
        score_table.rows[2].cells[1].text = f"{company.z_score:.2f}" if company.z_score else "N/A"
        score_table.rows[2].cells[2].text = "ALERT" if company.z_score and company.z_score < 1.8 else "OK"

        score_table.rows[3].cells[0].text = "F-Score"
        score_table.rows[3].cells[1].text = f"{int(company.f_score)}" if company.f_score else "N/A"
        score_table.rows[3].cells[2].text = "ALERT" if company.f_score and company.f_score < 3 else "OK"

        score_table.rows[4].cells[0].text = "Sloan Ratio"
        score_table.rows[4].cells[1].text = f"{company.sloan_ratio:.1f}%" if company.sloan_ratio else "N/A"
        score_table.rows[4].cells[2].text = "ALERT" if company.sloan_ratio and abs(company.sloan_ratio) > 10 else "OK"

        doc.add_paragraph()

        # AI Analysis
        if company.ai_summary:
            doc.add_heading("AI Analysis", level=3)
            doc.add_paragraph(company.ai_summary)

        if company.ai_recommendation:
            doc.add_heading("Recommendation", level=3)
            doc.add_paragraph(company.ai_recommendation)

        doc.add_paragraph("_" * 50)


def create_company_analysis_from_data(
    stock_data: Dict,
    anomaly_report: Optional[Any] = None,
    market: str = "US"
) -> CompanyAnalysis:
    """
    Create a CompanyAnalysis object from stock data and anomaly report.

    Args:
        stock_data: Dictionary of stock fundamental data
        anomaly_report: Optional anomaly detection report
        market: Market (US or SG)

    Returns:
        CompanyAnalysis object
    """
    def safe_get(key, default=None):
        val = stock_data.get(key, default)
        if pd.isna(val):
            return default
        return val

    analysis = CompanyAnalysis(
        symbol=safe_get('Symbol', 'N/A'),
        company_name=safe_get('Company', safe_get('Symbol', 'Unknown')),
        sector=safe_get('Sector', 'N/A'),
        market=market,

        gross_margin=safe_get('Gross Margin %'),
        net_margin=safe_get('Net Margin %'),
        roe=safe_get('ROE %'),
        roa=safe_get('ROA %'),
        debt_to_equity=safe_get('Debt-to-Equity'),
        fcf_margin=safe_get('FCF Margin %'),
        roic_wacc=safe_get('ROIC-WACC'),
        rote_wacc=safe_get('ROTE-WACC'),
        revenue_growth_5y=safe_get('5-Year Revenue Growth Rate (Per Share)'),
        eps_growth_5y=safe_get('5-Year EPS without NRI Growth Rate'),

        market_cap=safe_get('Market Cap ($M)'),
        epv=safe_get('Earnings Power Value (EPV)'),
        epv_mc_ratio=safe_get('EPV/MC Ratio'),
        valuation_status=safe_get('Valuation', 'N/A'),
    )

    # Add anomaly data if available
    if anomaly_report:
        analysis.risk_level = getattr(anomaly_report, 'risk_level', 'N/A')
        analysis.m_score = getattr(anomaly_report, 'm_score', None)
        analysis.z_score = getattr(anomaly_report, 'z_score', None)
        analysis.f_score = getattr(anomaly_report, 'f_score', None)
        analysis.sloan_ratio = getattr(anomaly_report, 'sloan_ratio', None)

        if hasattr(anomaly_report, 'anomalies'):
            analysis.anomalies = [
                {
                    'category': a.category,
                    'severity': a.severity.value,
                    'description': a.description,
                    'year': a.year,
                    'details': a.details
                }
                for a in anomaly_report.anomalies
            ]

    return analysis
