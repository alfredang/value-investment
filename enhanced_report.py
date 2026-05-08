"""
Enhanced AI-Powered Report Generator for Value Investment Academy
Generates professional investment reports with deep AI analysis
"""

import os
from io import BytesIO
from datetime import datetime
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


from llm import claude_complete


# Haiku keeps per-company cost low — the report makes 4-7 LLM calls per company.
REPORT_MODEL = "claude-haiku-4-5-20251001"

ANALYST_SYSTEM_PROMPT = """You are a senior equity research analyst at a top investment bank.
Write professional, insightful investment analysis. Be specific, data-driven,
and provide actionable insights. Use clear, concise financial language.
Do NOT use markdown formatting - write plain text suitable for a Word document."""


def get_anthropic_client(api_key: str = None):
    """No-op shim kept for backward compatibility with the public function signature.

    Auth is now handled by Claude Code CLI subscription login, not an API key.
    """
    return None


def generate_ai_analysis(client, prompt: str, max_tokens: int = 1500) -> str:
    """Generate AI analysis via claude-agent-sdk (uses Claude Code CLI auth).

    The `client` and `max_tokens` args are kept for signature stability;
    they are not used by claude-agent-sdk's query() under the hood.
    """
    try:
        return claude_complete(user=prompt, system=ANALYST_SYSTEM_PROMPT, model=REPORT_MODEL)
    except Exception as e:
        return f"[Analysis generation failed: {str(e)}]"


def generate_company_background(client, company_data: Dict) -> str:
    """Generate a 2-paragraph factual company background for the Profile section."""
    symbol = company_data.get('symbol', 'Unknown')
    company = company_data.get('company', 'Unknown Company')
    sector = company_data.get('sector', 'N/A')
    industry = company_data.get('industry', 'N/A')

    prompt = f"""Write a clean, professional 2-paragraph company background for inclusion in a research report.

Company: {company} (ticker: {symbol})
Sector: {sector}
Industry: {industry}

Output rules — strict:
- Output ONLY the two paragraphs of background prose. No preamble, no caveats about
  your training data, no "I have limited knowledge" disclaimers, no source URLs,
  no "Paragraph 1:" or "Paragraph 2:" labels, no markdown.
- Paragraph 1: What the company does (core products/services, customer base).
- Paragraph 2: Where it operates and its competitive position in its industry.
- 3-4 sentences per paragraph. Plain text suitable for a Word document.
- If you do not know the specific company well, write a generic, factual paragraph
  describing what businesses in this Sector/Industry typically do, framed as
  "{company} operates in the {industry} industry, where companies typically..." —
  do NOT meta-comment about your knowledge limits.

Begin the response with the first paragraph directly."""

    text = generate_ai_analysis(client, prompt, max_tokens=500)
    return _clean_background_text(text)


def _clean_background_text(text: str) -> str:
    """Strip any meta-commentary or labels the AI may emit despite instructions."""
    if not text:
        return text

    import re
    # Drop common preamble patterns the model sometimes opens with
    preamble_patterns = [
        r"^I (?:appreciate|understand|acknowledge|have|don't|cannot)[^.]*\.\s*",
        r"^Based on what I can reliably[^.]*[.:]\s*",
        r"^Let me (?:search|provide|write)[^.]*[.:]\s*",
        r"^Here (?:is|are)[^:]*:\s*",
        r"^Note:[^.]*\.\s*",
    ]
    cleaned = text.strip()
    for _ in range(3):  # repeat to catch stacked preambles
        for pat in preamble_patterns:
            cleaned = re.sub(pat, "", cleaned, flags=re.IGNORECASE).strip()

    # Drop literal paragraph labels like "PARAGRAPH 1:" / "Paragraph 1." / "**Paragraph 1**"
    cleaned = re.sub(r"\*{0,2}PARAGRAPH\s*\d+\s*\*{0,2}\s*[:.\-]\s*", "",
                     cleaned, flags=re.IGNORECASE)

    # Drop trailing CAVEAT/Sources/Disclaimer blocks
    for stop_word in ["CAVEAT:", "Sources:", "Source:", "Disclaimer:", "Note that",
                      "Given the time gap"]:
        idx = cleaned.find(stop_word)
        if idx > 0:
            cleaned = cleaned[:idx].rstrip()

    # Collapse 3+ blank lines into a single blank line
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()

    return cleaned


def compute_valuation_range(company_data: Dict) -> Dict[str, Any]:
    """
    Build the in-house Valuation Range using the client's two-stage DCF
    formula. Low IV and High IV are pre-computed in the screener; here we
    compose them into a 3-point range with a fair-value midpoint and verdict.

    Returns dict with keys: lower (Low IV), fair (midpoint), upper (High IV),
    current_price, currency, upside_pct, verdict.
    """
    low_iv = company_data.get('low_iv')
    high_iv = company_data.get('high_iv')
    current_price = company_data.get('current_price')
    currency = company_data.get('currency', 'USD')

    out = {
        'lower': None,
        'fair': None,
        'upper': None,
        'current_price': current_price,
        'currency': currency,
        'upside_pct': None,
        'verdict': 'N/A',
    }

    def _num(v):
        try:
            f = float(v) if v is not None else None
            return f if f is None or f == f else None
        except (TypeError, ValueError):
            return None

    low = _num(low_iv)
    high = _num(high_iv)
    cp = _num(current_price)

    # Fall back to EPV-based range only if the in-house formula couldn't run
    if low is None or high is None or low <= 0 or high <= 0:
        epv_val = _num(company_data.get('epv'))
        if epv_val is None or epv_val <= 0:
            out['verdict'] = 'N/A (Negative EPS — DCF math undefined)'
            return out
        out['lower'] = round(epv_val * 0.7, 2)
        out['fair'] = round(epv_val, 2)
        out['upper'] = round(epv_val * 1.3, 2)
    else:
        out['lower'] = round(low, 2)
        out['fair'] = round((low + high) / 2.0, 2)
        out['upper'] = round(high, 2)

    if cp is not None and cp > 0 and out['fair'] is not None:
        out['upside_pct'] = round(((out['fair'] - cp) / cp) * 100, 1)
        if cp < out['lower']:
            out['verdict'] = 'Undervalued'
        elif cp <= out['upper']:
            out['verdict'] = 'Fair Value'
        else:
            out['verdict'] = 'Overvalued'

    return out


def generate_executive_summary(client, report_data: List[Dict], criteria: Dict) -> str:
    """Generate AI-powered executive summary."""
    undervalued = [d for d in report_data if d.get('valuation') == 'Undervalued']
    clean_companies = [d for d in report_data if d.get('ai_rating') == 'CLEAN']
    minor_companies = [d for d in report_data if d.get('ai_rating') == 'MINOR']
    material_companies = [d for d in report_data if d.get('ai_rating') == 'MATERIAL']

    # Calculate average metrics
    avg_roe = sum(d.get('roe', 0) or 0 for d in report_data) / max(len(report_data), 1)
    avg_margin = sum(d.get('gross_margin', 0) or 0 for d in report_data) / max(len(report_data), 1)

    symbols = [d['symbol'] for d in undervalued[:5]]

    prompt = f"""Write a compelling executive summary (3-4 paragraphs) for a value investment research report.

PORTFOLIO OVERVIEW:
- Total companies analyzed: {len(report_data)}
- Undervalued opportunities (EPV/MC > 1.3): {len(undervalued)} companies
- Clean financials (no anomalies): {len(clean_companies)} companies
- Minor anomalies detected: {len(minor_companies)} companies
- Material distortions found: {len(material_companies)} companies
- Average ROE across universe: {avg_roe:.1f}%
- Average Gross Margin: {avg_margin:.1f}%

TOP UNDERVALUED PICKS: {', '.join(symbols) if symbols else 'None identified'}

SCREENING CRITERIA USED:
- Minimum Gross Margin: {criteria.get('gross_margin', 'N/A')}%
- Minimum Net Margin: {criteria.get('net_margin', 'N/A')}%
- Minimum ROE: {criteria.get('roe', 'N/A')}%
- Maximum Debt-to-Equity: {criteria.get('debt_to_equity', 'N/A')}

Write a professional executive summary that:
1. Opens with the key investment thesis and market opportunity
2. Highlights the most compelling investment candidates with specific reasoning
3. Summarizes the risk-reward profile of the screened universe
4. Provides clear actionable recommendations for investors

Write in a sophisticated, institutional investor style. Be specific about the value proposition."""

    return generate_ai_analysis(client, prompt)


def generate_company_deep_dive(client, company_data: Dict) -> Dict[str, str]:
    """Generate comprehensive AI analysis for a single company."""

    symbol = company_data.get('symbol', 'Unknown')
    company = company_data.get('company', 'Unknown Company')

    # Extract all available metrics
    metrics = {
        'valuation': company_data.get('valuation', 'N/A'),
        'epv': company_data.get('epv'),
        'market_cap': company_data.get('market_cap'),
        'ai_rating': company_data.get('ai_rating', 'N/A'),
        'ai_analysis': company_data.get('ai_analysis', ''),
        'roe': company_data.get('roe'),
        'gross_margin': company_data.get('gross_margin'),
        'net_margin': company_data.get('net_margin'),
        'debt_equity': company_data.get('debt_equity'),
        'fcf_margin': company_data.get('fcf_margin'),
        'roic_wacc': company_data.get('roic_wacc'),
        'rev_growth': company_data.get('rev_growth'),
        'eps_growth': company_data.get('eps_growth')
    }

    # Calculate EPV/MC ratio if available
    epv_mc_ratio = None
    margin_of_safety = None
    if metrics['epv'] and metrics['market_cap']:
        try:
            epv_mc_ratio = float(metrics['epv']) / float(metrics['market_cap'])
            margin_of_safety = (epv_mc_ratio - 1) * 100
        except:
            pass

    # Risk flags from AI anomaly analysis
    ai_rating = metrics.get('ai_rating', 'N/A')
    anomaly_risk = "LOW" if ai_rating == 'CLEAN' else "MODERATE" if ai_rating == 'MINOR' else "ELEVATED"

    # Business Analysis
    business_prompt = f"""Provide a concise business analysis (2-3 paragraphs) for {symbol} ({company}).

FINANCIAL PROFILE:
- Return on Equity: {metrics['roe']}%
- Gross Margin: {metrics['gross_margin']}%
- Net Margin: {metrics['net_margin']}%
- Debt-to-Equity: {metrics['debt_equity']}
- Free Cash Flow Margin: {metrics['fcf_margin']}%
- 5Y Revenue Growth: {metrics['rev_growth']}%
- 5Y EPS Growth: {metrics['eps_growth']}%

Based on these metrics, analyze:
1. The quality and sustainability of the business model
2. Competitive positioning and moat indicators
3. Capital allocation efficiency and management quality signals

Write as a senior analyst providing insights to portfolio managers. Be specific about what the numbers tell us."""

    business_analysis = generate_ai_analysis(client, business_prompt, max_tokens=800)

    # Risk Assessment
    ai_analysis_excerpt = metrics.get('ai_analysis', '')[:500] if metrics.get('ai_analysis') else 'No AI anomaly analysis available.'

    risk_prompt = f"""Provide a comprehensive risk assessment (2-3 paragraphs) for {symbol}.

RISK INDICATORS:
- AI Anomaly Rating: {ai_rating} (Anomaly Risk: {anomaly_risk})
- Debt-to-Equity: {metrics['debt_equity']}
- ROIC-WACC Spread: {metrics.get('roic_wacc', 'N/A')}

AI ANOMALY ANALYSIS FINDINGS:
{ai_analysis_excerpt}

Analyze:
1. What the AI anomaly findings tell us about earnings quality and financial consistency
2. Specific red flags or concerns investors should monitor
3. Mitigating factors or reasons for confidence

Be balanced but highlight genuine concerns."""

    risk_analysis = generate_ai_analysis(client, risk_prompt, max_tokens=700)

    # Investment Thesis
    thesis_prompt = f"""Write a clear investment thesis (2 paragraphs) for {symbol} ({company}).

KEY FACTS:
- Valuation: {metrics['valuation']} (EPV/MC: {f"{epv_mc_ratio:.2f}" if epv_mc_ratio else 'N/A'})
- ROE: {metrics['roe']}% | Gross Margin: {metrics['gross_margin']}%
- AI Anomaly Rating: {ai_rating} ({anomaly_risk} risk)
- Growth: Revenue {metrics['rev_growth']}% | EPS {metrics['eps_growth']}%

Write a compelling investment thesis that:
1. Summarizes the bull case in clear terms
2. Identifies key risks and monitoring points
3. Provides a clear recommendation (Strong Buy / Buy / Hold / Avoid)

Be decisive and specific about why an investor should or should not own this stock."""

    investment_thesis = generate_ai_analysis(client, thesis_prompt, max_tokens=600)

    return {
        'business_analysis': business_analysis,
        'risk_analysis': risk_analysis,
        'investment_thesis': investment_thesis,
    }


def generate_ceo_analysis(client, report_data: List[Dict]) -> str:
    """Generate AI analysis of management quality indicators."""

    # Aggregate management quality signals from the data
    high_roe_companies = [d for d in report_data if d.get('roe') and d['roe'] > 15]
    clean_financials = [d for d in report_data if d.get('ai_rating') == 'CLEAN']
    low_debt = [d for d in report_data if d.get('debt_equity') and d['debt_equity'] < 0.5]

    prompt = f"""Write a professional analysis (3-4 paragraphs) of management quality indicators across the analyzed companies.

DATA POINTS:
- Companies with ROE > 15% (indicates capital discipline): {len(high_roe_companies)}
- Companies with clean financials (no anomalies): {len(clean_financials)}
- Companies with D/E < 0.5 (indicates conservative financing): {len(low_debt)}
- Total companies analyzed: {len(report_data)}

Top performers by ROE: {', '.join([d['symbol'] for d in sorted(report_data, key=lambda x: x.get('roe') or 0, reverse=True)[:3]])}
Companies with cleanest financials: {', '.join([d['symbol'] for d in clean_financials[:3]]) if clean_financials else 'None'}

Provide analysis on:
1. What these quantitative signals tell us about management quality across the portfolio
2. Which companies show the strongest evidence of shareholder-aligned management
3. Red flags that might indicate poor capital allocation or governance
4. Recommendations for further qualitative due diligence

Write as a governance analyst providing institutional investors with management quality insights."""

    return generate_ai_analysis(client, prompt, max_tokens=1000)


def generate_portfolio_recommendations(client, report_data: List[Dict]) -> str:
    """Generate AI-powered portfolio construction recommendations."""

    # Rank companies by combined score
    ranked = []
    for d in report_data:
        score = 0
        if d.get('valuation') == 'Undervalued':
            score += 3
        elif d.get('valuation') == 'Fair':
            score += 1
        # AI anomaly rating scoring
        ai_rating = d.get('ai_rating', '')
        if ai_rating == 'CLEAN':
            score += 3
        elif ai_rating == 'MINOR':
            score += 1
        if d.get('roe') and d['roe'] > 15:
            score += 2
        elif d.get('roe') and d['roe'] > 10:
            score += 1
        if d.get('roic_wacc') and d['roic_wacc'] > 5:
            score += 1
        ranked.append({**d, 'composite_score': score})

    ranked.sort(key=lambda x: x['composite_score'], reverse=True)

    top_picks = ranked[:5]
    top_picks_summary = "\n".join([
        f"- {d['symbol']}: Score {d['composite_score']}/10, {d.get('valuation', 'N/A')}, ROE {d.get('roe')}%, AI Rating: {d.get('ai_rating', 'N/A')}"
        for d in top_picks
    ])

    avoid_list = [d for d in ranked if d['composite_score'] <= 2]
    avoid_summary = ", ".join([d['symbol'] for d in avoid_list[:5]]) if avoid_list else "None"

    prompt = f"""Write professional portfolio recommendations (4-5 paragraphs) for a value-focused portfolio.

TOP RANKED PICKS (Composite Score out of 10):
{top_picks_summary}

COMPANIES TO AVOID OR MONITOR: {avoid_summary}

UNIVERSE STATISTICS:
- Total companies: {len(report_data)}
- Strong conviction picks (Score ≥ 7): {len([d for d in ranked if d['composite_score'] >= 7])}
- Moderate conviction (Score 4-6): {len([d for d in ranked if 4 <= d['composite_score'] < 7])}
- Low conviction (Score < 4): {len([d for d in ranked if d['composite_score'] < 4])}

Write recommendations covering:
1. Model portfolio construction with specific allocations for top picks
2. Sector/concentration considerations
3. Risk management and position sizing guidance
4. Monitoring triggers and rebalancing criteria
5. Timeline for position building

Write as a portfolio strategist advising a value-focused investment committee."""

    return generate_ai_analysis(client, prompt, max_tokens=1200)


def set_document_styles(doc: Document):
    """Configure professional document styling."""

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue

    # Heading 2 style
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 76, 153)

    # Heading 3 style
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Calibri'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(51, 102, 153)


def add_styled_table(doc: Document, headers: List[str], rows: List[List[str]],
                     header_color: RGBColor = RGBColor(0, 51, 102)):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header
        # Make header bold
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        # Set header background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '003366')  # Dark blue
        cell._tc.get_or_add_tcPr().append(shading)
        # Set header text color to white
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Add data rows with zebra striping for readability
    zebra_fill = 'F4F6F9'  # very light blue-gray
    for r_idx, row_data in enumerate(rows):
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = str(value)
            if r_idx % 2 == 1:  # alternate rows
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), zebra_fill)
                row[i]._tc.get_or_add_tcPr().append(shading)

    return table


def fmt_score(val, decimals=2):
    """Format a numeric score."""
    return f"{val:.{decimals}f}" if val is not None else "N/A"


def fmt_pct(val):
    """Format a percentage."""
    return f"{val}%" if val is not None else "N/A"


def fmt_money(val):
    """Format money in millions."""
    return f"${val}M" if val is not None else "N/A"


def _parse_chart_json(raw: str) -> Optional[Dict[str, list]]:
    """Parse a JSON blob with {symbol, years: [{year, revenue_m, net_income_m, fcf_m}]}
    into the chart engine's series dict format. Returns None if invalid."""
    import json
    import re

    if not raw:
        return None
    cleaned = raw.strip()
    # Strip code fences if Claude wrapped the JSON despite instructions
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"```\s*$", "", cleaned).strip()
    # Defensive: find outermost JSON object
    first = cleaned.find("{")
    last = cleaned.rfind("}")
    if first == -1 or last == -1 or last <= first:
        return None
    cleaned = cleaned[first:last + 1]
    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError:
        return None

    years = data.get("years") or []
    if not years:
        return None

    rev_pts, ni_pts, fcf_pts = [], [], []
    for entry in years:
        y = str(entry.get("year", "")).strip()
        if not y:
            continue
        for arr, key in [(rev_pts, "revenue_m"), (ni_pts, "net_income_m"), (fcf_pts, "fcf_m")]:
            v = entry.get(key)
            if v is not None:
                try:
                    arr.append((y, float(v)))
                except (TypeError, ValueError):
                    pass

    series: Dict[str, list] = {}
    if rev_pts:
        series["Revenue ($M)"] = rev_pts
    if ni_pts:
        series["Net Income ($M)"] = ni_pts
    if fcf_pts:
        series["Free Cash Flow ($M)"] = fcf_pts
    return series or None


def _fetch_ten_year_via_firecrawl(
    symbol: str, company: str, market: str = "US"
) -> Optional[Dict[str, list]]:
    """
    Scrape REAL 10-year financials via Firecrawl (stockanalysis.com), then
    have Claude extract the numbers from the scraped tables — NOT from
    Claude's training memory.

    Returns chart-engine series dict, or None if Firecrawl failed.
    """
    from llm import claude_complete
    try:
        from scraper import scrape_financial_data, is_firecrawl_available
    except ImportError:
        return None

    if not is_firecrawl_available():
        return None

    scraped = scrape_financial_data(symbol, company or symbol, market)
    if not scraped.get("ok"):
        return None

    md = scraped.get("markdown", "")
    if not md:
        return None

    # Cap to ~25k chars to keep prompts light — financial tables are usually
    # near the top of these pages anyway.
    md_for_prompt = md[:25000]

    system = (
        "You are a strict data-extraction tool. Read the provided real financial "
        "data scraped from public filings and return JSON ONLY — no preamble, no "
        "commentary, no markdown code fences. Use null for any year you cannot "
        "find. Convert local currency to USD millions when obvious; otherwise "
        "keep the source currency and note it."
    )
    user = f"""From the SCRAPED financial data below, extract annual Revenue, Net Income, and Free Cash Flow for the past 5-10 fiscal years for {company} ({symbol}).

DATA SOURCE (REAL, scraped via Firecrawl):
{md_for_prompt}

Return ONLY this exact JSON shape (no other text, no fences):

{{
  "symbol": "{symbol}",
  "currency": "USD or BRL or whatever the source uses — match the table",
  "years": [
    {{"year": "2020", "revenue_m": 1201, "net_income_m": 308, "fcf_m": 234}},
    {{"year": "2021", "revenue_m": 1753, "net_income_m": 440, "fcf_m": 461}}
  ]
}}

Rules:
- Order years OLDEST first
- Use null for any number you cannot find in the scraped data
- Don't invent numbers — only use what's explicitly in the data
- If FCF isn't listed, derive it as Operating Cash Flow minus CapEx if both are present, otherwise null"""

    try:
        raw = claude_complete(user=user, system=system)
    except Exception:
        return None

    return _parse_chart_json(raw)


def _fetch_ten_year_via_claude(symbol: str, company: str) -> Optional[Dict[str, list]]:
    """
    LAST-RESORT fallback when Firecrawl fails: ask Claude to estimate from
    its training memory. Output is approximate and the caption flags this
    clearly.
    """
    from llm import claude_complete

    system = (
        "You are a financial-data assistant. Output STRICT JSON only — no preamble, "
        "no commentary, no markdown code fences, no caveats. If you genuinely do not "
        "have data for a company, return {\"years\": []}."
    )
    user = f"""Provide your best estimate of {company} ({symbol})'s annual financial data for the past 10 fiscal years based on your training knowledge.

For each year, give:
- Total Revenue (in USD millions, converted from local currency if needed)
- Net Income (in USD millions)
- Free Cash Flow (in USD millions)

Return ONLY this exact JSON shape, no other text:

{{
  "symbol": "{symbol}",
  "years": [
    {{"year": "2015", "revenue_m": 25.0, "net_income_m": 2.0, "fcf_m": 3.5}},
    {{"year": "2016", "revenue_m": 35.0, "net_income_m": 3.0, "fcf_m": 4.0}}
  ]
}}

Order years oldest first. Use null for any field you cannot estimate.
Output ONLY the JSON."""

    try:
        raw = claude_complete(user=user, system=system)
    except Exception:
        return None
    return _parse_chart_json(raw)


def _synthesize_history_from_csv_row(company_data: Dict) -> Optional[Dict[str, list]]:
    """
    Build an approximate 5-year history from the screener CSV metrics, for
    cases where Claude has no training data on the company. Computes:
      Net Income (current)  = Market Cap / PE Ratio
      Revenue (current)     = Net Income / Net Margin
      FCF (current)         = Revenue × FCF Margin
    Then back-extrapolates using the 5-year growth rates already in the CSV.

    Returns None if the inputs are insufficient (e.g., loss-making company
    with no PE ratio).
    """
    from datetime import datetime

    def _to_float(v):
        try:
            f = float(v)
            return f if f == f else None  # filter NaN
        except (TypeError, ValueError):
            return None

    market_cap = _to_float(company_data.get('market_cap'))
    pe_ratio = _to_float(company_data.get('pe_ratio'))
    net_margin = _to_float(company_data.get('net_margin'))
    fcf_margin = _to_float(company_data.get('fcf_margin'))
    rev_growth = _to_float(company_data.get('rev_growth'))
    eps_growth = _to_float(company_data.get('eps_growth'))
    fcf_growth = _to_float(company_data.get('fcf_growth'))

    if not market_cap or market_cap <= 0:
        return None
    if not pe_ratio or pe_ratio <= 0:
        return None
    if not net_margin or net_margin <= 0:
        # Negative-margin / loss-making company — math doesn't apply
        return None

    current_ni = market_cap / pe_ratio
    current_revenue = current_ni / (net_margin / 100.0)
    current_fcf = current_revenue * (fcf_margin / 100.0) if fcf_margin else None

    # Default growth = 0% if missing
    rev_factor = 1 + ((rev_growth or 0) / 100.0)
    eps_factor = 1 + ((eps_growth or 0) / 100.0)
    fcf_factor = 1 + ((fcf_growth or rev_growth or 0) / 100.0)

    # 10-year window: last completed fiscal year + 9 years prior
    current_year = datetime.now().year - 1
    years = list(range(current_year - 9, current_year + 1))

    revenue_series, ni_series, fcf_series = [], [], []
    for i, year in enumerate(years):
        years_back = (len(years) - 1) - i
        if rev_factor > 0:
            rev = current_revenue / (rev_factor ** years_back)
            revenue_series.append((str(year), round(rev, 1)))
        if eps_factor > 0:
            ni = current_ni / (eps_factor ** years_back)
            ni_series.append((str(year), round(ni, 1)))
        if current_fcf is not None and fcf_factor > 0:
            fcf = current_fcf / (fcf_factor ** years_back)
            fcf_series.append((str(year), round(fcf, 1)))

    series = {}
    if revenue_series:
        series["Revenue ($M)"] = revenue_series
    if ni_series:
        series["Net Income ($M)"] = ni_series
    if fcf_series:
        series["Free Cash Flow ($M)"] = fcf_series
    return series or None


def _embed_ten_year_chart(doc, symbol: str, company: str = "", company_data: Optional[Dict] = None) -> bool:
    """
    Embed a 5-10 year line chart in the DOCX.

    Strategy (waterfall):
      1. PREFERRED: Firecrawl scrapes real filings from stockanalysis.com,
         then Claude extracts the numbers from those real tables.
      2. CSV-based synthesis (math-derived extrapolation from screening metrics).
      3. LAST RESORT: Claude estimates from training memory (clearly flagged).
      4. If all three fail, return False so caller prints a placeholder message.
    """
    try:
        from chart_engine import make_ten_year_line_chart_png

        # Determine market for Firecrawl URL routing
        market = "US"
        if company_data:
            ex = (company_data.get('exchange') or '').upper()
            if 'SGX' in ex:
                market = "SG"

        source_label = ""
        series = None

        # Step 1: Firecrawl — REAL data scraped from public filings
        try:
            series = _fetch_ten_year_via_firecrawl(symbol, company or symbol, market)
            if series:
                source_label = (
                    f"Source: stockanalysis.com (scraped via Firecrawl) — annual filings for {symbol}. "
                    "Numbers extracted directly from public filings tables."
                )
        except Exception:
            series = None

        # Step 2: CSV synthesis fallback (math-derived from screening metrics)
        if not series and company_data:
            series = _synthesize_history_from_csv_row(company_data)
            if series:
                source_label = (
                    f"Source: Model-derived 10-year extrapolation from current metrics + 5-year growth "
                    f"rates (Market Cap, PE, Margins). Indicative only — cross-check against 10-K "
                    f"filings before any investment decision."
                )

        # Step 3: Claude training memory (LAST resort, clearly flagged)
        if not series:
            series = _fetch_ten_year_via_claude(symbol, company or symbol)
            if series:
                source_label = (
                    f"Source: Claude AI estimates from training memory — figures for {symbol} are "
                    f"approximate. Verify against 10-K / primary filings before any investment decision."
                )

        if not series:
            return False

        png_buf = make_ten_year_line_chart_png(
            title=f"{symbol} — 10-Year Revenue, Net Income, and Free Cash Flow",
            series_by_label=series,
            y_format="money",
        )
        if png_buf is None:
            return False

        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_run = pic_para.add_run()
        pic_run.add_picture(png_buf, width=Inches(6.0))

        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(source_label)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(120, 120, 120)
        return True
    except Exception as e:
        ph = doc.add_paragraph()
        ph_run = ph.add_run(f"[10-year chart generation failed for {symbol}: {e}]")
        ph_run.italic = True
        ph_run.font.color.rgb = RGBColor(180, 80, 80)
        return True


def generate_professional_report(
    report_data: List[Dict],
    criteria: Dict,
    api_key: str,
    progress_callback=None,
    universe_df=None,
) -> BytesIO:
    """
    Generate a professional AI-enhanced investment report.

    Args:
        report_data: List of company data dictionaries
        criteria: Screening criteria used
        api_key: Ignored. Auth via Claude Code CLI login. Kept for signature stability.
        progress_callback: Optional callback function for progress updates

    Returns:
        BytesIO buffer containing the DOCX file
    """

    client = None  # No client object needed; claude-agent-sdk handles auth via CLI
    doc = Document()

    # Apply professional styling
    set_document_styles(doc)

    def update_progress(message: str):
        if progress_callback:
            progress_callback(message)

    # ========================================
    # COVER PAGE
    # ========================================
    update_progress("Creating cover page...")

    # Vertical spacing to push content down toward visual center
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_heading('VIA FINANCIAL ANALYSIS PLATFORM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Professional Investment Research Report', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tagline.add_run("EPV Valuation  ·  Forensic Risk Assessment  ·  AI-Enhanced Analysis")
    tag_run.italic = True
    tag_run.font.size = Pt(11)
    tag_run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    # Metadata block
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_color = RGBColor(80, 80, 80)

    def _meta_line(label: str, value: str, bold_value: bool = False):
        run_lab = meta.add_run(f"{label}: ")
        run_lab.font.size = Pt(11)
        run_lab.font.color.rgb = label_color
        run_val = meta.add_run(f"{value}\n")
        run_val.font.size = Pt(11)
        run_val.bold = bold_value

    _meta_line("Report Date", datetime.now().strftime('%B %d, %Y'), bold_value=True)
    _meta_line("Generated", datetime.now().strftime('%Y-%m-%d %H:%M'))
    _meta_line("Companies Analyzed", str(len(report_data)))
    _meta_line("Analysis Framework", "EPV Valuation + Forensic Risk Assessment")

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Confidentiality footer
    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cf_run = confidential.add_run("CONFIDENTIAL — For internal review only")
    cf_run.font.size = Pt(9)
    cf_run.font.color.rgb = RGBColor(120, 120, 120)
    cf_run.bold = True

    doc.add_page_break()

    # ========================================
    # TABLE OF CONTENTS
    # ========================================
    doc.add_heading('Table of Contents', level=1)

    toc_items = [
        ("1.", "Executive Summary", "AI-generated investment thesis and key findings"),
        ("2.", "Screening Results Overview", "Quantitative screening results and valuation distribution"),
        ("3.", "Detailed Company Analysis",
            "Per-company sections: A. Profile  ·  B. Valuation Range  ·  "
            "C. Key Metrics  ·  D. 10-Year Trends  ·  E. Competitor Comparison  ·  F. Analyst Assessment"),
        ("4.", "Management Quality Assessment", "Governance indicators and capital allocation discipline"),
        ("5.", "Portfolio Recommendations", "Position sizing and allocation strategy"),
        ("6.", "Risk Disclosure & Methodology", "Methodology notes, limitations, and disclaimer"),
    ]

    for num, title, desc in toc_items:
        p = doc.add_paragraph()
        p.add_run(f"{num} {title}").bold = True
        p.add_run(f"\n    {desc}")

    doc.add_page_break()

    # ========================================
    # 1. EXECUTIVE SUMMARY (AI-Generated)
    # ========================================
    update_progress("Generating AI executive summary...")

    doc.add_heading('1. Executive Summary', level=1)

    executive_summary = generate_executive_summary(client, report_data, criteria)

    # Add the AI-generated summary
    for paragraph in executive_summary.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    doc.add_paragraph("")

    # Key metrics summary table
    doc.add_heading('Key Metrics at a Glance', level=2)

    undervalued = [d for d in report_data if d.get('valuation') == 'Undervalued']
    clean = [d for d in report_data if d.get('ai_rating') == 'CLEAN']
    minor = [d for d in report_data if d.get('ai_rating') == 'MINOR']
    material = [d for d in report_data if d.get('ai_rating') == 'MATERIAL']

    add_styled_table(doc,
        ['Metric', 'Count', 'Percentage'],
        [
            ['Total Companies Screened', str(len(report_data)), '100%'],
            ['Undervalued (EPV/MC > 1.3)', str(len(undervalued)), f'{len(undervalued)*100//max(len(report_data),1)}%'],
            ['Clean Financials (No Anomalies)', str(len(clean)), f'{len(clean)*100//max(len(report_data),1)}%'],
            ['Minor Anomalies Detected', str(len(minor)), f'{len(minor)*100//max(len(report_data),1)}%'],
            ['Material Distortions Found', str(len(material)), f'{len(material)*100//max(len(report_data),1)}%']
        ]
    )

    doc.add_page_break()

    # ========================================
    # 2. SCREENING RESULTS OVERVIEW
    # ========================================
    update_progress("Building screening results section...")

    doc.add_heading('2. Screening Results Overview', level=1)

    doc.add_heading('Screening Criteria Applied', level=2)
    criteria_table = add_styled_table(doc,
        ['Criterion', 'Threshold', 'Description'],
        [
            ['Gross Margin', f"≥ {criteria.get('gross_margin', 'N/A')}%", 'Pricing power indicator'],
            ['Net Margin', f"≥ {criteria.get('net_margin', 'N/A')}%", 'Profitability measure'],
            ['Return on Equity', f"≥ {criteria.get('roe', 'N/A')}%", 'Capital efficiency'],
            ['Debt-to-Equity', f"≤ {criteria.get('debt_to_equity', 'N/A')}", 'Financial leverage'],
            ['FCF Margin', f"≥ {criteria.get('fcf_margin', 'N/A')}%", 'Cash generation ability'],
            ['ROIC - WACC', f"≥ {criteria.get('roic_wacc', 'N/A')}%", 'Value creation spread']
        ]
    )

    doc.add_paragraph("")

    doc.add_heading('Valuation Distribution', level=2)

    fair_value = [d for d in report_data if d.get('valuation') == 'Fair']
    overvalued = [d for d in report_data if d.get('valuation') == 'Overvalued']

    add_styled_table(doc,
        ['Valuation Category', 'Count', 'Companies'],
        [
            ['UNDERVALUED (EPV/MC > 1.3)', str(len(undervalued)),
             ', '.join([d['symbol'] for d in undervalued][:6]) + ('...' if len(undervalued) > 6 else '')],
            ['FAIR VALUE (0.7 ≤ EPV/MC ≤ 1.3)', str(len(fair_value)),
             ', '.join([d['symbol'] for d in fair_value][:6]) + ('...' if len(fair_value) > 6 else '')],
            ['OVERVALUED (EPV/MC < 0.7)', str(len(overvalued)),
             ', '.join([d['symbol'] for d in overvalued][:6]) + ('...' if len(overvalued) > 6 else '')]
        ]
    )

    doc.add_paragraph("")

    doc.add_heading('Complete Results Summary', level=2)

    # Prepare rows for results table
    results_rows = []
    for d in report_data:
        results_rows.append([
            d.get('symbol', 'N/A'),
            str(d.get('company', 'N/A'))[:20],
            d.get('valuation', 'N/A'),
            d.get('ai_rating', 'N/A'),
            fmt_pct(d.get('roe')),
            fmt_pct(d.get('gross_margin')),
            fmt_pct(d.get('fcf_margin'))
        ])

    add_styled_table(doc,
        ['Symbol', 'Company', 'Valuation', 'AI Rating', 'ROE', 'Gross Margin', 'FCF Margin'],
        results_rows
    )

    doc.add_page_break()

    # ========================================
    # 3. DETAILED COMPANY ANALYSIS (Standardized Template)
    # ========================================
    doc.add_heading('3. Detailed Company Analysis', level=1)

    doc.add_paragraph(
        "Each company in this section follows a standardized research template: "
        "Company Profile, Valuation Range, Key Financial Metrics, 10-Year Financial "
        "Trends, Competitor Comparison, and Analyst Assessment."
    )

    for idx, company_data in enumerate(report_data):
        symbol = company_data.get('symbol', 'Unknown')
        company = company_data.get('company', 'Unknown Company')

        update_progress(f"Analyzing {symbol} ({idx + 1}/{len(report_data)})...")

        # Section header
        doc.add_heading(f"3.{idx + 1} {symbol} — {company}", level=2)

        # Generate AI analysis for this company (background + risk + thesis)
        background = generate_company_background(client, company_data)
        ai_analysis = generate_company_deep_dive(client, company_data)
        valuation = compute_valuation_range(company_data)
        currency = valuation.get('currency') or company_data.get('currency', 'USD')

        # ----------------------------------------------------------------
        # A. COMPANY PROFILE
        # ----------------------------------------------------------------
        doc.add_heading('A. Company Profile', level=3)

        add_styled_table(doc,
            ['Field', 'Value'],
            [
                ['Company Name', company],
                ['Ticker Symbol', symbol],
                ['Exchange', str(company_data.get('exchange', 'N/A'))],
                ['Sector', str(company_data.get('sector', 'N/A'))],
                ['Industry', str(company_data.get('industry', 'N/A'))],
                ['Sub-Industry', str(company_data.get('subindustry', 'N/A'))],
                ['Reporting Currency', currency],
            ]
        )
        doc.add_paragraph("")

        # Background paragraphs
        for para in background.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())

        doc.add_paragraph("")

        # ----------------------------------------------------------------
        # B. VALUATION RANGE
        # ----------------------------------------------------------------
        doc.add_heading('B. Valuation Range', level=3)

        cp_str = f"{currency} {valuation['current_price']:.2f}" if valuation.get('current_price') is not None else 'N/A'
        upside_str = f"{valuation['upside_pct']:+.1f}% vs current price" if valuation.get('upside_pct') is not None else 'N/A'

        add_styled_table(doc,
            ['Scenario', 'Price per Share', 'Description'],
            [
                ['Low Intrinsic Value (Conservative)',
                 f"{currency} {valuation['lower']:.2f}" if valuation.get('lower') is not None else 'N/A',
                 'In-house DCF — historical EPS growth dampened by 8× (capped at 3%)'],
                ['Fair Value (midpoint)',
                 f"{currency} {valuation['fair']:.2f}" if valuation.get('fair') is not None else 'N/A',
                 'Midpoint of the Low/High IV range'],
                ['High Intrinsic Value (Aggressive)',
                 f"{currency} {valuation['upper']:.2f}" if valuation.get('upper') is not None else 'N/A',
                 'In-house DCF — historical EPS growth dampened by 2× (capped at 12%)'],
                ['Current Market Price', cp_str, 'Market quote at time of analysis'],
                ['Implied Upside to Fair Value', upside_str, 'Positive = trading below midpoint'],
                ['Verdict', valuation.get('verdict', 'N/A'),
                 'Undervalued < Low IV; Fair Value within range; Overvalued > High IV'],
            ]
        )
        doc.add_paragraph("")
        doc.add_paragraph(
            "Valuation derived from the client's in-house two-stage DCF formula: "
            "EPS growth dampened to a conservative scenario (Low IV) and an aggressive "
            "scenario (High IV), each discounted at 10% with a 2% terminal growth rate "
            "and a 30% margin of safety. Inputs sourced from the screener CSV "
            "(Current Price, PE Ratio, 5-Year EPS Growth)."
        )
        doc.add_paragraph("")

        # ----------------------------------------------------------------
        # C. KEY FINANCIAL METRICS
        # ----------------------------------------------------------------
        doc.add_heading('C. Key Financial Metrics', level=3)
        add_styled_table(doc,
            ['Metric', 'Value', 'Notes'],
            [
                ['Return on Equity (ROE)', fmt_pct(company_data.get('roe')), 'Profit per $1 of shareholder equity'],
                ['Return on Assets (ROA)', fmt_pct(company_data.get('roa')), 'Profit per $1 of total assets'],
                ['Gross Margin', fmt_pct(company_data.get('gross_margin')), 'Pricing power signal'],
                ['Net Margin', fmt_pct(company_data.get('net_margin')), 'Bottom-line profitability'],
                ['FCF Margin', fmt_pct(company_data.get('fcf_margin')), 'Cash generation per $1 of revenue'],
                ['Debt-to-Equity', str(company_data.get('debt_equity', 'N/A')), 'Balance sheet leverage'],
                ['ROIC − WACC', fmt_pct(company_data.get('roic_wacc')), 'Value creation spread'],
                ['ROTE − WACC', fmt_pct(company_data.get('rote_wacc')), 'Value creation on tangible equity'],
                ['5-Year Revenue Growth', fmt_pct(company_data.get('rev_growth')), 'Topline momentum'],
                ['5-Year EPS Growth', fmt_pct(company_data.get('eps_growth')), 'Earnings momentum per share'],
                ['Market Capitalization', fmt_money(company_data.get('market_cap')), 'Total equity value'],
            ]
        )
        doc.add_paragraph("")

        # ----------------------------------------------------------------
        # D. 10-YEAR FINANCIAL TRENDS (line charts via FMP API)
        # ----------------------------------------------------------------
        doc.add_heading('D. 10-Year Financial Trends', level=3)

        history_added = _embed_ten_year_chart(doc, symbol, company, company_data=company_data)
        if not history_added:
            ph = doc.add_paragraph()
            ph_run = ph.add_run(
                f"[Insufficient data to construct a historical chart for {symbol}: "
                "either the company is loss-making (no positive PE ratio), or the "
                "screener CSV is missing key fields (Market Cap, Net Margin, Growth Rates).]"
            )
            ph_run.italic = True
            ph_run.font.color.rgb = RGBColor(120, 120, 120)
        doc.add_paragraph("")

        # ----------------------------------------------------------------
        # E. COMPETITOR COMPARISON (bar charts)
        # ----------------------------------------------------------------
        doc.add_heading('E. Competitor Comparison', level=3)

        peer_charts_added = False
        if universe_df is not None and not universe_df.empty:
            try:
                from peer_finder import find_peers, build_peer_metrics_frame
                from chart_engine import make_competitor_bar_chart_png

                target_rows = universe_df[universe_df['Symbol'] == symbol]
                if not target_rows.empty:
                    target_row_full = target_rows.iloc[0]
                    peers = find_peers(symbol, universe_df, limit=5)

                    if not peers.empty:
                        intro = doc.add_paragraph()
                        intro.add_run(
                            f"{symbol} is compared below against {len(peers)} peer{'s' if len(peers) != 1 else ''} "
                            f"({', '.join(peers['Symbol'].tolist())}) sharing the same Sector and Industry, "
                            "ranked by market capitalization."
                        )

                        metrics_df = build_peer_metrics_frame(target_row_full, peers)
                        peer_chart_specs = [
                            ('ROE %', 'Return on Equity', True),
                            ('Net Margin %', 'Net Margin', True),
                            ('FCF Margin %', 'Free Cash Flow Margin', True),
                            ('5-Year Revenue Growth Rate (Per Share)', '5-Year Revenue Growth', True),
                            ('Debt-to-Equity', 'Debt-to-Equity Ratio', False),
                        ]

                        for col_name, label, is_pct in peer_chart_specs:
                            png_buf = make_competitor_bar_chart_png(
                                metric_label=label,
                                metrics_df=metrics_df,
                                metric_column=col_name,
                                target_symbol=symbol,
                                is_percentage=is_pct,
                            )
                            if png_buf is not None:
                                pic_para = doc.add_paragraph()
                                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                pic_run = pic_para.add_run()
                                pic_run.add_picture(png_buf, width=Inches(6.0))
                                peer_charts_added = True
                    else:
                        ph = doc.add_paragraph()
                        ph_run = ph.add_run(
                            f"[No peers found for {symbol} in the screening universe — "
                            "widen the screen criteria to include more companies in the same Sector + Industry.]"
                        )
                        ph_run.italic = True
                        ph_run.font.color.rgb = RGBColor(120, 120, 120)
                        peer_charts_added = True  # explanatory text counts
            except Exception as e:
                ph = doc.add_paragraph()
                ph_run = ph.add_run(f"[Competitor chart generation failed: {e}]")
                ph_run.italic = True
                ph_run.font.color.rgb = RGBColor(180, 80, 80)
                peer_charts_added = True

        if not peer_charts_added:
            ph = doc.add_paragraph()
            ph_run = ph.add_run(
                f"[Bar charts comparing {symbol} against peers in the {company_data.get('industry', 'same')} "
                "industry require the screening universe — re-run Step 1 before generating the report.]"
            )
            ph_run.italic = True
            ph_run.font.color.rgb = RGBColor(120, 120, 120)

        doc.add_paragraph("")

        # ----------------------------------------------------------------
        # F. ANALYST ASSESSMENT
        # ----------------------------------------------------------------
        doc.add_heading('F. Analyst Assessment', level=3)

        # F.1 Business Quality
        doc.add_heading('Business Quality', level=4)
        for para in ai_analysis['business_analysis'].split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())

        # F.2 Risk Assessment
        doc.add_heading('Risk Assessment', level=4)

        ai_rating = company_data.get('ai_rating', 'N/A')
        rating_desc = {
            'CLEAN': 'No significant distortions detected',
            'MINOR': 'Minor one-offs detected, not material to valuation',
            'MATERIAL': 'Material distortions found, EPV adjustment may be needed',
        }
        add_styled_table(doc,
            ['Indicator', 'Reading', 'Interpretation'],
            [
                ['AI Anomaly Rating', ai_rating, rating_desc.get(ai_rating, 'Analysis not available')],
                ['Debt-to-Equity', str(company_data.get('debt_equity', 'N/A')),
                 'Balance sheet leverage indicator'],
                ['ROIC-WACC Spread', fmt_pct(company_data.get('roic_wacc')),
                 'Value creation above cost of capital'],
            ]
        )
        doc.add_paragraph("")

        ai_anomaly_text = company_data.get('ai_analysis', '')
        if ai_anomaly_text:
            doc.add_paragraph("AI anomaly findings:")
            for para in ai_anomaly_text.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.add_paragraph("")

        for para in ai_analysis['risk_analysis'].split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())

        # F.3 Investment Thesis
        doc.add_heading('Investment Thesis', level=4)
        for para in ai_analysis['investment_thesis'].split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())

        doc.add_page_break()

    # ========================================
    # 4. MANAGEMENT QUALITY ASSESSMENT
    # ========================================
    update_progress("Generating management quality assessment...")

    doc.add_heading('4. Management Quality Assessment', level=1)

    doc.add_paragraph(
        "Management quality is a critical determinant of long-term shareholder returns. "
        "While quantitative metrics provide signals, qualitative assessment of leadership "
        "commitment and capital allocation discipline is essential for value investors."
    )

    doc.add_heading('Quantitative Signals of Management Quality', level=2)

    # Generate AI management analysis
    ceo_analysis = generate_ceo_analysis(client, report_data)

    for para in ceo_analysis.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_paragraph("")

    # Management quality ranking table
    doc.add_heading('Management Quality Indicators by Company', level=2)

    mgmt_rows = []
    for d in sorted(report_data, key=lambda x: (x.get('roe') or 0), reverse=True):
        roe = d.get('roe')
        ai_rating = d.get('ai_rating', 'N/A')

        if ai_rating == 'CLEAN' and roe and roe > 15:
            quality = "HIGH"
            signal = "Strong execution, clean financials"
        elif ai_rating in ('CLEAN', 'MINOR') and roe and roe > 10:
            quality = "ABOVE AVG"
            signal = "Good capital discipline"
        elif ai_rating == 'MATERIAL':
            quality = "CAUTION"
            signal = "Material distortions, requires scrutiny"
        else:
            quality = "AVERAGE"
            signal = "Monitor execution"

        mgmt_rows.append([
            d.get('symbol', 'N/A'),
            quality,
            ai_rating,
            fmt_pct(roe),
            signal
        ])

    add_styled_table(doc,
        ['Symbol', 'Quality Rating', 'AI Rating', 'ROE', 'Key Signal'],
        mgmt_rows
    )

    doc.add_paragraph("")

    doc.add_heading('Due Diligence Recommendations', level=2)
    doc.add_paragraph("For each company, we recommend reviewing:")
    doc.add_paragraph("• Proxy statements for executive compensation alignment")
    doc.add_paragraph("• Insider trading patterns over the past 12 months")
    doc.add_paragraph("• Management tenure and track record of guidance accuracy")
    doc.add_paragraph("• Capital allocation history (M&A returns, buyback timing, dividend policy)")
    doc.add_paragraph("• Related party transactions and governance structure")

    doc.add_page_break()

    # ========================================
    # 5. PORTFOLIO RECOMMENDATIONS
    # ========================================
    update_progress("Generating portfolio recommendations...")

    doc.add_heading('5. Portfolio Recommendations', level=1)

    # Generate AI portfolio recommendations
    portfolio_recs = generate_portfolio_recommendations(client, report_data)

    for para in portfolio_recs.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_paragraph("")

    # Final recommendation table
    doc.add_heading('Summary Recommendations', level=2)

    rec_rows = []
    for d in report_data:
        ai_clean = d.get('ai_rating') == 'CLEAN'
        ai_minor = d.get('ai_rating') == 'MINOR'
        ai_ok = ai_clean or ai_minor
        underval = d.get('valuation') == 'Undervalued'
        strong_roe = bool(d.get('roe') and d['roe'] > 15)

        score = sum([ai_clean, ai_ok, underval, strong_roe])

        if underval and ai_clean and strong_roe:
            rating = "STRONG BUY"
            action = "Initiate 3-5% position"
        elif underval and ai_ok:
            rating = "BUY"
            action = "Initiate 2-3% position"
        elif d.get('valuation') == 'Fair' and ai_ok:
            rating = "HOLD"
            action = "Monitor for entry"
        else:
            rating = "WATCH"
            action = "Further research needed"

        rec_rows.append([
            d.get('symbol', 'N/A'),
            d.get('valuation', 'N/A'),
            rating,
            f"{score}/4",
            action
        ])

    # Sort by conviction
    rec_rows.sort(key=lambda x: (
        0 if 'STRONG' in x[2] else 1 if 'BUY' == x[2] else 2 if 'HOLD' in x[2] else 3
    ))

    add_styled_table(doc,
        ['Symbol', 'Valuation', 'Rating', 'Conviction', 'Suggested Action'],
        rec_rows
    )

    doc.add_page_break()

    # ========================================
    # 6. RISK DISCLOSURE
    # ========================================
    doc.add_heading('6. Risk Disclosure & Methodology', level=1)

    doc.add_heading('Analysis Methodology', level=2)

    doc.add_paragraph("This report employs a multi-factor value investing framework:")
    doc.add_paragraph("")

    add_styled_table(doc,
        ['Metric', 'Purpose', 'Threshold', 'Interpretation'],
        [
            ['EPV/Market Cap', 'Intrinsic value assessment', '> 1.3', 'Undervalued with margin of safety'],
            ['AI Anomaly Rating', 'Detect one-off financial distortions', 'CLEAN', 'No material distortions in 10Y history'],
            ['ROIC - WACC', 'Value creation spread', '> 0', 'Creating value above cost of capital'],
            ['ROE', 'Capital efficiency', '> 15%', 'Strong returns on equity']
        ]
    )

    doc.add_paragraph("")

    doc.add_heading('Important Limitations', level=2)
    doc.add_paragraph("• Analysis is based on historical financial data and may not reflect current conditions")
    doc.add_paragraph("• AI-generated insights are supplementary and should not replace professional judgment")
    doc.add_paragraph("• Quantitative screens may miss qualitative factors critical to investment success")
    doc.add_paragraph("• Market conditions can change rapidly, affecting valuations and risk profiles")
    doc.add_paragraph("• Past performance and financial metrics do not guarantee future results")

    doc.add_paragraph("")

    doc.add_heading('Disclaimer', level=2)

    disclaimer_para = doc.add_paragraph()
    disclaimer_para.add_run("IMPORTANT: ").bold = True
    disclaimer_para.add_run(
        "This report is for informational and educational purposes only and does not constitute "
        "investment advice, an offer to sell, or a solicitation of an offer to buy any securities. "
        "The analysis and recommendations contained herein are based on publicly available information "
        "and AI-assisted analysis, which may contain errors or omissions. Investors should conduct "
        "their own due diligence and consult with qualified financial advisors before making any "
        "investment decisions. Past performance is not indicative of future results. All investments "
        "involve risk, including the possible loss of principal."
    )

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("─" * 40)
    footer.add_run(f"\nGenerated by VIA Financial Analysis Platform | {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    footer.add_run("\nPowered by AI-Enhanced Analysis")

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    update_progress("Report generation complete!")

    return buffer
